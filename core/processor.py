import pandas as pd
import random
import os
from typing import List, Dict, Any, Tuple, Optional
from fpdf import FPDF
import logging
from datetime import datetime
import textwrap
import re
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.shared import qn
import tempfile
import math2docx
from .template_generator import create_test_template
from .neural_query_generator import create_neural_query_document

def format_number_with_comma(number: float, decimals: int = 1) -> str:
    """Форматирует число с запятой вместо точки как десятичный разделитель"""
    formatted = f"{number:.{decimals}f}"
    return formatted.replace('.', ',')


def get_task_type_indicator(question: Dict[str, Any]) -> str:
    """Повертає індикатор типу завдання для питання
    
    Args:
        question: Словник з даними питання
        
    Returns:
        Рядок з індикатором типу завдання
    """
    if question.get('is_test_question', True):
        # Тестове питання - перевіряємо кількість правильних відповідей
        correct_answer = str(question.get('correct_answer', '')).strip().upper()
        # Перевіряємо чи є кілька літер у відповіді (множинний вибір)
        if len(correct_answer) > 1 or (isinstance(question.get('correct_answer', ''), str) and ',' in question.get('correct_answer', '')):
            return "(Виберіть декілька правильних відповідей)"
        else:
            return "(Виберіть одну правильну відповідь)"
    else:
        # Відкрите питання
        return "(Запишіть відповідь)"


def process_math_formulas(text: str) -> str:
    """Обробляє математичні формули в тексті та конвертує їх у LaTeX формат
    
    Args:
        text: Текст з математичними формулами
        
    Returns:
        Текст з обробленими формулами у LaTeX форматі
    """
    if not text:
        return text
    
    processed_text = text
    
    # Спочатку обробляємо спеціальні символи □
    processed_text = processed_text.replace('□', '')
    
    # Паттерни для розпізнавання математичних виразів (в порядку від складних до простих)
    patterns = [
        # Складні вкладені дроби: \frac{\frac{a}{b}}{c} або \frac{a}{\frac{b}{c}}
        (r'\\frac\{\\frac\{([^}]+)\}\{([^}]+)\}\}\{([^}]+)\}', r'\\frac{\\frac{\1}{\2}}{\3}'),
        (r'\\frac\{([^}]+)\}\{\\frac\{([^}]+)\}\{([^}]+)\}\}', r'\\frac{\1}{\\frac{\2}{\3}}'),
        
        # Дроби з дужками: (чисельник)/(знаменник)
        (r'\(([^)]+)\)/\(([^)]+)\)', r'\\frac{\1}{\2}'),
        
        # Прості дроби: число/число або змінна/змінна
        (r'([a-zA-Z0-9]+)/([a-zA-Z0-9]+)', r'\\frac{\1}{\2}'),
        
        # Степені з дужками: (вираз)^показник
        (r'\(([^)]+)\)\^([a-zA-Z0-9]+)', r'(\1)^{\2}'),
        
        # Степені: основа^показник
        (r'([a-zA-Z0-9]+)\^([a-zA-Z0-9]+)', r'\1^{\2}'),
        
        # Корені: sqrt(вираз)
        (r'sqrt\(([^)]+)\)', r'\\sqrt{\1}'),
    ]
    
    # Застосовуємо паттерни послідовно
    for pattern, replacement in patterns:
        processed_text = re.sub(pattern, replacement, processed_text)
    
    # Видаляємо зайві дужки навколо простих виразів
    # Паттерн для видалення дужок навколо простих дробів
    processed_text = re.sub(r'\(\\frac\{([^}]+)\}\{([^}]+)\}\)', r'\\frac{\1}{\2}', processed_text)
    
    # Видаляємо дужки навколо одиночних змінних або чисел
    processed_text = re.sub(r'\(([a-zA-Z0-9])\)', r'\1', processed_text)
    
    return processed_text


def add_formatted_text_to_paragraph(paragraph, text: str):
    """Додає текст з математичними формулами до параграфа Word документа
    
    Args:
        paragraph: Параграф Word документа
        text: Текст з можливими математичними формулами
    """
    # Обробляємо текст для виявлення математичних формул
    processed_text = process_math_formulas(text)
    
    # Розділяємо текст на частини: звичайний текст та математичні формули
    # Паттерн для пошуку LaTeX формул (включаючи вкладені дроби)
    math_pattern = r'(\\frac\{(?:[^{}]|\\frac\{[^}]+\}\{[^}]+\})+\}\{(?:[^{}]|\\frac\{[^}]+\}\{[^}]+\})+\}|[a-zA-Z0-9()]+\^\{[^}]+\}|\\sqrt\{[^}]+\})'
    
    # Розділяємо текст на частини
    parts = re.split(math_pattern, processed_text)
    
    for part in parts:
        if not part:  # Пропускаємо порожні частини
            continue
            
        # Перевіряємо, чи це математична формула
        if (part.startswith('\\frac{') or 
            ('^{' in part and any(c.isalnum() for c in part)) or 
            part.startswith('\\sqrt{')):
            try:
                # Створюємо справжню математичну формулу через math2docx
                math2docx.add_math(paragraph, part)
            except Exception as e:
                # Якщо не вдалося створити формулу, додаємо курсивом
                run = paragraph.add_run(part)
                run.italic = True
        else:
            # Звичайний текст
            paragraph.add_run(part)

log = logging.getLogger(__name__)

def ensure_temp_dir(prefix="temp_"):
    """Создает временную папку внутри проекта"""
    try:
        # Получаем корневую папку проекта
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        temp_dir = os.path.join(project_root, "temp")
        
        # Создаем папку temp если её нет
        os.makedirs(temp_dir, exist_ok=True)
        
        # Создаем подпапку с префиксом
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        specific_temp_dir = os.path.join(temp_dir, f"{prefix}{timestamp}")
        os.makedirs(specific_temp_dir, exist_ok=True)
        
        return specific_temp_dir
    except Exception as e:
        log.warning(f"Не удалось создать временную папку: {e}. Используем системную временную папку.")
        return tempfile.gettempdir()

def get_text_width(pdf: FPDF, text: str) -> float:
    """Получить ширину текста в текущем шрифте"""
    return pdf.get_string_width(text)

def fit_text_to_width(pdf: FPDF, text: str, max_width: float, min_font_size: int = 8) -> Tuple[str, int]:
    """Подогнать текст под заданную ширину, уменьшая шрифт или разбивая на строки"""
    current_font = pdf.font_family
    current_style = pdf.font_style
    current_size = pdf.font_size_pt
    
    # Попробуем уменьшить шрифт
    for font_size in range(int(current_size), min_font_size - 1, -1):
        pdf.set_font(current_font, current_style, font_size)
        if get_text_width(pdf, text) <= max_width:
            return text, font_size
    
    # Если не помогло уменьшение шрифта, разбиваем текст
    pdf.set_font(current_font, current_style, min_font_size)
    words = text.split()
    lines = []
    current_line = ""
    
    for word in words:
        test_line = current_line + (" " if current_line else "") + word
        if get_text_width(pdf, test_line) <= max_width:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
                current_line = word
            else:
                # Слово слишком длинное, принудительно разбиваем
                lines.append(word)
    
    if current_line:
        lines.append(current_line)
    
    return "\n".join(lines), min_font_size

def add_multiline_text(pdf: FPDF, text: str, max_width: float, line_height: float = 6, min_font_size: int = 8):
    """Добавить многострочный текст с автоматическим переносом"""
    fitted_text, font_size = fit_text_to_width(pdf, text, max_width, min_font_size)
    
    current_font = pdf.font_family
    current_style = pdf.font_style
    pdf.set_font(current_font, current_style, font_size)
    
    lines = fitted_text.split('\n')
    for line in lines:
        pdf.cell(0, line_height, line, ln=True)
    
    return len(lines) * line_height

def read_test_excel(file_path: str) -> pd.DataFrame:
    """
    Читає Excel файл з питаннями тесту.
    
    Очікувана структура:
    - Рядок 1: Інструкції (пропускається)
    - Стовпець 0: Номер питання
    - Стовпець 1: Текст питання
    - Стовпець 2: Номер правильної відповіді (або порожній для нетестових завдань)
    - Стовпець 3: Вага завдання (за замовчуванням 1)
    - Стовпці 4+: Варіанти відповідей (опціонально для нетестових завдань)
    
    Args:
        file_path: Шлях до Excel файлу
        
    Returns:
        DataFrame з питаннями тесту
    """
    try:
        # Читаємо Excel файл, пропускаючи перший рядок з інструкціями
        df = pd.read_excel(file_path, header=None, skiprows=1)
        
        # Конвертуємо стовпці з номером питання та текстом питання у рядки
        # Інші стовпці залишаємо як є для збереження початкового форматування
        df[df.columns[0]] = df[df.columns[0]].astype(str)  # Номер вопроса
        df[df.columns[1]] = df[df.columns[1]].astype(str)  # Текст вопроса
        
        # Проверяем минимальную структуру (номер + вопрос + правильный ответ + вес + варианты + тип)
        if df.shape[1] < 6:
            raise ValueError("Файл повинен містити мінімум 6 стовпців: номер питання, питання, правильна відповідь, вага завдання, варіант А, тип завдання")
        
        # Удаляем пустые строки (проверяем наличие номера вопроса и текста вопроса)
        df = df[df.iloc[:, 0].notna() & (df.iloc[:, 0] != 'nan') & 
               df.iloc[:, 1].notna() & (df.iloc[:, 1] != 'nan')]  # Видаляємо рядки де немає номера або питання
        
        if df.empty:
            raise ValueError("Файл не містить валідних даних")
        
        # Перейменовуємо стовпці для зручності - адаптуємося до фактичної кількості колонок
        base_columns = ['question_number', 'question', 'correct_answer', 'weight']
        # Варіанти відповідей: загальна кількість колонок мінус базові колонки (4) мінус тип завдання (1)
        option_columns = [f'option_{i+1}' for i in range(max(0, df.shape[1] - 5))]  # Варіанти відповідей
        task_type_column = ['task_type']
        
        columns = base_columns + option_columns + task_type_column
        
        # Обрізаємо список колонок до фактичної кількості
        columns = columns[:df.shape[1]]
        df.columns = columns
        
        # Додаємо відсутні колонки якщо потрібно
        required_columns = ['question_number', 'question', 'correct_answer', 'weight', 'option_1', 'option_2', 'option_3', 'option_4', 'task_type']
        for col in required_columns:
            if col not in df.columns:
                df[col] = '' if col != 'weight' else 1.0
        
        # Обробляємо вагу завдання (за замовчуванням 1)
        df['weight'] = pd.to_numeric(df['weight'], errors='coerce')
        df['weight'] = df['weight'].fillna(1.0)  # Заповнюємо порожні значення одиницею
        
        # Визначаємо тип завдання на основі колонки task_type
        df['task_type'] = df['task_type'].astype(str).str.strip()
        
        # Підраховуємо кількість непорожніх варіантів відповідей для кожного питання
        option_cols = ['option_1', 'option_2', 'option_3', 'option_4']
        df['option_count'] = 0
        df['has_gaps'] = False
        
        # Перевіряємо послідовність варіантів відповідей для кожного рядка
        for idx in df.index:
            task_type = df.loc[idx, 'task_type']
            consecutive_count = 0
            has_gap = False
            found_empty = False
            has_first_option = False
            
            # Для відкритих питань варіанти не потрібні
            if task_type.lower() in ['відкрите', 'відкрите питання', 'вп']:
                df.at[idx, 'option_count'] = 0
                df.at[idx, 'has_gaps'] = False
                continue
            
            for i, col in enumerate(option_cols):
                cell_value = df.loc[idx, col]
                is_empty = pd.isna(cell_value) or cell_value == 'nan' or str(cell_value).strip() == ''
                
                if i == 0:  # Перший варіант
                    has_first_option = not is_empty
                
                if not is_empty:
                    if found_empty:  # Знайшли заповнений варіант після порожнього
                        has_gap = True
                        break
                    consecutive_count += 1
                else:
                    found_empty = True
            
            # Якщо є будь-які варіанти відповідей, але перший порожній - це помилка
            has_any_options = any(not (pd.isna(df.loc[idx, col]) or df.loc[idx, col] == 'nan' or str(df.loc[idx, col]).strip() == '') for col in option_cols)
            if has_any_options and not has_first_option:
                has_gap = True
            
            df.at[idx, 'option_count'] = consecutive_count
            df.at[idx, 'has_gaps'] = has_gap
        
        # Перевіряємо наявність пропусків у варіантах відповідей для тестових питань
        test_questions_mask = ~df['task_type'].str.lower().isin(['відкрите', 'відкрите питання', 'вп'])
        gaps_mask = df['has_gaps'] & test_questions_mask
        if gaps_mask.any():
            gap_questions = df[gaps_mask][['question_number', 'question']].values.tolist()
            error_details = []
            for q_num, q_text in gap_questions[:5]:  # Показуємо перші 5 помилок
                error_details.append(f"Питання {q_num}: '{q_text[:50]}...'")
            error_msg = f"Знайдено тестові питання з пропусками у варіантах відповідей. Варіанти повинні йти підряд без пропусків, починаючи з першого варіанту:\n" + "\n".join(error_details)
            if len(gap_questions) > 5:
                error_msg += f"\n... та ще {len(gap_questions) - 5} питань"
            raise ValueError(error_msg)
        
        # Перевіряємо обов'язкове заповнення правильної відповіді
        missing_answers = df['correct_answer'].isna() | (df['correct_answer'] == 'nan') | (df['correct_answer'].astype(str).str.strip() == '')
        if missing_answers.any():
            missing_questions = df[missing_answers][['question_number', 'question']].values.tolist()
            error_details = []
            for q_num, q_text in missing_questions[:5]:  # Показуємо перші 5 помилок
                error_details.append(f"Питання {q_num}: '{q_text[:50]}...'")
            error_msg = f"Знайдено питання без правильної відповіді:\n" + "\n".join(error_details)
            if len(missing_questions) > 5:
                error_msg += f"\n... та ще {len(missing_questions) - 5} питань"
            error_msg += "\n\nВсі питання повинні мати правильну відповідь."
            raise ValueError(error_msg)
        
        # Перевіряємо тестові питання з недостатньою кількістю варіантів відповіді
        test_questions_mask = ~df['task_type'].str.lower().isin(['відкрите', 'відкрите питання', 'вп'])
        single_option_mask = (df['option_count'] == 1) & test_questions_mask
        if single_option_mask.any():
            single_option_questions = df[single_option_mask][['question_number', 'question']].values.tolist()
            error_details = []
            for q_num, q_text in single_option_questions[:5]:  # Показуємо перші 5 помилок
                error_details.append(f"Питання {q_num}: '{q_text[:50]}...'")
            error_msg = f"Знайдено тестові питання з одним варіантом відповіді. Для тестових питань має бути мінімум два варіанти:\n" + "\n".join(error_details)
            if len(single_option_questions) > 5:
                error_msg += f"\n... та ще {len(single_option_questions) - 5} питань"
            raise ValueError(error_msg)
        
        # Визначаємо тип завдання на основі колонки task_type
        df['is_test_question'] = ~df['task_type'].str.lower().isin(['відкрите', 'відкрите питання', 'вп']) & df['correct_answer'].notna() & (df['correct_answer'] != 'nan')
        df['is_multiple_choice'] = df['task_type'].str.lower().isin(['тест м', 'тестове з декількома варіантами відповіді', 'тестове завдання з декількома варіантами відповіді', 'тк'])
        
        # Строга валідація типів завдань - всі питання повинні мати розпізнаний тип
        valid_task_types = [
            'відкрите', 'відкрите питання', 'вп',  # Відкриті завдання
            'тест', 'тестове', 'тестове завдання', 'то',  # Тестові завдання з одним варіантом
            'тест м', 'тестове з декількома варіантами відповіді', 'тестове завдання з декількома варіантами відповіді', 'тк'  # Тестові завдання з кількома варіантами
        ]
        
        # Перевіряємо чи всі типи завдань розпізнані
        unrecognized_mask = ~df['task_type'].str.lower().isin(valid_task_types)
        if unrecognized_mask.any():
            unrecognized_questions = df[unrecognized_mask][['question_number', 'question', 'task_type']].values.tolist()
            error_details = []
            for q_num, q_text, task_type in unrecognized_questions[:5]:  # Показуємо перші 5 помилок
                error_details.append(f"Питання {q_num}: '{q_text[:50]}...' - невідомий тип завдання: '{task_type}'")
            
            valid_types_str = "\n".join([
                "Відкриті завдання: 'відкрите', 'відкрите питання', 'вп'",
                "Тестові завдання (один варіант): 'тест', 'тестове', 'тестове завдання', 'то'",
                "Тестові завдання (кілька варіантів): 'тест м', 'тестове з декількома варіантами відповіді', 'тестове завдання з декількома варіантами відповіді', 'тк'"
            ])
            
            error_msg = f"Знайдено питання з невизначеними типами завдань:\n" + "\n".join(error_details)
            if len(unrecognized_questions) > 5:
                error_msg += f"\n... та ще {len(unrecognized_questions) - 5} питань"
            error_msg += f"\n\nДоступні типи завдань:\n{valid_types_str}"
            raise ValueError(error_msg)
        
        # Приводимо стовпець correct_answer до object типу для уникнення попереджень
        df['correct_answer'] = df['correct_answer'].astype('object')
        
        # Для тестових завдань перевіряємо правильні відповіді
        test_mask = df['is_test_question']
        if test_mask.any():
            # Українські літери (без Ґ, Є, І, Ї, Й, Ь)
            ukrainian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ю', 'Я']
            
            # Перевіряємо правильні відповіді для тестових завдань
            for idx in df[test_mask].index:
                answer_str = str(df.loc[idx, 'correct_answer']).strip().upper()
                option_count = df.loc[idx, 'option_count']
                q_num = df.loc[idx, 'question_number']
                q_text = df.loc[idx, 'question']
                
                # Перетворюємо числові індекси в українські літери
                converted_answer = ""
                for char in answer_str:
                    if char.isdigit():
                        # Перетворюємо число в українську літеру
                        digit_idx = int(char) - 1  # -1 оскільки нумерація з 1
                        if 0 <= digit_idx < len(ukrainian_letters):
                            converted_answer += ukrainian_letters[digit_idx]
                        else:
                            raise ValueError(f"Питання {q_num}: '{q_text[:50]}...' - некоректний індекс '{char}' в відповіді '{answer_str}'. Доступні індекси: 1-{len(ukrainian_letters)}")
                    elif char in ukrainian_letters:
                        converted_answer += char
                    else:
                        raise ValueError(f"Питання {q_num}: '{q_text[:50]}...' - некоректний символ '{char}' в відповіді '{answer_str}'. Доступні літери: {', '.join(ukrainian_letters)} або числа 1-{len(ukrainian_letters)}")
                
                # Перевіряємо чи відповідь не пуста
                if not converted_answer:
                    raise ValueError(f"Питання {q_num}: '{q_text[:50]}...' - порожня відповідь")
                
                # Зберігаємо перетворену відповідь
                df.at[idx, 'correct_answer'] = converted_answer
        
        # Для відкритих завдань зберігаємо правильну відповідь як текст
        open_mask = ~df['is_test_question'] & df['correct_answer'].notna() & (df['correct_answer'] != 'nan')
        if open_mask.any():
            # Для відкритих завдань форматуємо відповідь правильно
            for idx in df[open_mask].index:
                answer_value = df.loc[idx, 'correct_answer']
                # Якщо це число, форматуємо без зайвих .0
                if isinstance(answer_value, (int, float)) and answer_value == int(answer_value):
                    df.at[idx, 'correct_answer'] = str(int(answer_value))
                else:
                    df.at[idx, 'correct_answer'] = str(answer_value).strip()
        
        # Видаляємо тимчасовий стовпець
        df = df.drop('option_count', axis=1)
        
        # Приводимо стовпець correct_answer до строкового типу для сумісності з PyArrow
        df['correct_answer'] = df['correct_answer'].astype(str)
        
        # Приводимо всі option_ колонки до строкового типу для сумісності з PyArrow
        option_cols = [col for col in df.columns if col.startswith('option_')]
        for col in option_cols:
            df[col] = df[col].astype(str)
        
        log.info(f"Завантажено {len(df)} питань з файлу {file_path}")
        return df
        
    except Exception as e:
        log.error(f"Помилка при читанні файлу {file_path}: {e}")
        raise

def _process_optional_questions(df: pd.DataFrame) -> pd.DataFrame:
    """
    Обробляє опціональні питання (коли номери повторюються).
    Вибирає по одному випадковому питанню для кожного номера.
    
    Args:
        df: DataFrame з питаннями
        
    Returns:
        DataFrame з обраними питаннями
    """
    # Групуємо питання за номерами
    question_groups = df.groupby('question_number')
    
    selected_questions = []
    
    for question_num, group in question_groups:
        if len(group) > 1:
            # Якщо є кілька питань з одним номером, вибираємо випадковий
            selected_question = group.sample(n=1).copy()
            log.info(f"Для номера питання {question_num} обрано випадковий варіант з {len(group)} доступних")
        else:
            # Якщо питання одне, просто беремо його
            selected_question = group.copy()
        
        selected_questions.append(selected_question)
    
    # Об'єднуємо обрані питання
    result_df = pd.concat(selected_questions, ignore_index=True)
    
    # Сортуємо за номером питання для збереження порядку
    result_df = result_df.sort_values('question_number').reset_index(drop=True)
    
    log.info(f"Оброблено {len(df)} питань, обрано {len(result_df)} унікальних")
    return result_df

def generate_test_variants(df: pd.DataFrame, num_variants: int, question_shuffle_mode: str = 'full', answer_shuffle_mode: str = 'random') -> List[Dict[str, Any]]:
    """
    Генерує варіанти тестів з перемішаними питаннями та відповідями.
    
    Args:
        df: DataFrame з питаннями
        num_variants: Кількість варіантів для генерації
        question_shuffle_mode: Режим перемішування питань ('full', 'easy_to_hard', 'none')
        answer_shuffle_mode: Режим перемішування варіантів відповідей ('random', 'none')
        
    Returns:
        Список словників з варіантами тестів
    """
    variants = []
    
    for variant_num in range(1, num_variants + 1):
        # Обробляємо опціональні питання для кожного варіанту окремо
        processed_df = _process_optional_questions(df)
        variant = {
            'variant_number': variant_num,
            'questions': [],
            'answer_key': []
        }
        
        # Упорядковуємо питання залежно від обраного режиму
        if question_shuffle_mode == 'full':
            # Повне перемішування
            shuffled_df = processed_df.sample(frac=1).reset_index(drop=True)
        elif question_shuffle_mode == 'easy_to_hard':
            # Сортування від легкого до складного (за вагою)
            shuffled_df = processed_df.sort_values('weight').reset_index(drop=True)
        else:  # question_shuffle_mode == 'none'
            # Не перемішуємо, залишаємо початковий порядок
            shuffled_df = processed_df.reset_index(drop=True)
        
        for idx, row in shuffled_df.iterrows():
            question_data = {
                'question_text': str(row['question']),
                'weight': float(row['weight']),
                'is_test_question': bool(row['is_test_question'])
            }
            
            if row['is_test_question']:
                # Тестове завдання з варіантами відповідей
                options = []
                for col in df.columns:
                    if col.startswith('option_') and pd.notna(row[col]) and str(row[col]).strip() != '' and str(row[col]) != 'nan':
                        # Зберігаємо початкове форматування чисел
                        value = row[col]
                        if isinstance(value, (int, float)):
                            # Для чисел: цілі без .0, дробові як є
                            if isinstance(value, float) and value.is_integer():
                                options.append(str(int(value)))
                            else:
                                options.append(str(value))
                        else:
                            options.append(str(value).strip())
                
                if len(options) < 2:
                    raise ValueError(f"Тестове питання '{row['question'][:50]}...' має менше 2 варіантів відповідей. Тестові питання повинні мати мінімум 2 варіанти відповіді.")
                
                # Перевіряємо коректність правильної відповіді
                correct_answer_str = str(row['correct_answer']).strip().upper()
                
                # Обробляємо множинні відповіді (наприклад, АВ, БГ, АВДЖИ)
                ukrainian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ю', 'Я']
                correct_answer_indices = []
                correct_option_texts = []
                
                # Перевіряємо кожну літеру в відповіді
                for letter in correct_answer_str:
                    if letter in ukrainian_letters[:len(options)]:
                        idx = ukrainian_letters.index(letter)
                        if idx < len(options):
                            correct_answer_indices.append(idx)
                            correct_option_texts.append(options[idx])
                    else:
                        # Спробуємо як число
                        try:
                            idx = int(letter) - 1
                            if 0 <= idx < len(options):
                                correct_answer_indices.append(idx)
                                correct_option_texts.append(options[idx])
                        except ValueError:
                            pass
                
                if not correct_answer_indices:
                    raise ValueError(f"Некоректна правильна відповідь '{row['correct_answer']}' для питання '{row['question'][:50]}...'. Правильна відповідь повинна містити літери А-Я або числа 1-{len(options)}.")
                
                # Перемішуємо варіанти відповідей залежно від режиму
                if answer_shuffle_mode == 'random':
                    shuffled_options = options.copy()
                    random.shuffle(shuffled_options)
                else:  # answer_shuffle_mode == 'none'
                    shuffled_options = options.copy()
                
                # Знаходимо нові позиції правильних відповідей після перемішування
                new_correct_letters = []
                for correct_text in correct_option_texts:
                    new_position_idx = shuffled_options.index(correct_text)
                    new_correct_letter = ukrainian_letters[new_position_idx] if new_position_idx < len(ukrainian_letters) else str(new_position_idx + 1)
                    new_correct_letters.append(new_correct_letter)
                
                # Об'єднуємо літери в одну відповідь
                combined_correct_answer = ''.join(sorted(new_correct_letters))
                
                question_data.update({
                    'options': shuffled_options,
                    'correct_answer': combined_correct_answer
                })
            else:
                # Відкрите завдання (без варіантів відповідей)
                # Обробляємо правильну відповідь як текстові дані
                formatted_answer = str(row['correct_answer']).strip()
                
                question_data.update({
                    'correct_answer': formatted_answer,
                    'options': []  # Відкриті питання не мають варіантів відповідей
                })
            
            variant['questions'].append(question_data)
            
            # Додаємо до ключа відповідей
            if row['is_test_question']:
                variant['answer_key'].append(combined_correct_answer)
            else:
                variant['answer_key'].append(formatted_answer)
        
        variants.append(variant)
        log.info(f"Згенеровано варіант {variant_num} з {len(variant['questions'])} питаннями")
    
    return variants

def create_test_pdf(variants: List[Dict[str, Any]], output_dir: str, columns: int = 1) -> Tuple[str, str]:
    """
    Створює PDF файли з тестами для учнів та відповідями для вчителя.
    
    Args:
        variants: Список варіантів тестів
        output_dir: Папка для збереження файлів
        columns: Кількість колонок для розміщення питань (1-3)
        
    Returns:
        Кортеж (шлях до файлу з тестами, шлях до файлу з відповідями)
    """
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    test_pdf_path = os.path.join(output_dir, f"tests_{timestamp}.pdf")
    answers_pdf_path = os.path.join(output_dir, f"answers_{timestamp}.pdf")
    
    # Створюємо PDF з тестами для учнів
    test_pdf = FPDF()
    test_pdf.add_font('Arial', '', 'c:/windows/fonts/arial.ttf', uni=True)
    test_pdf.add_font('Arial', 'B', 'c:/windows/fonts/arialbd.ttf', uni=True)
    
    page_width = test_pdf.w - 2 * test_pdf.l_margin  # Ширина сторінки без полів
    
    for variant in variants:
        test_pdf.add_page()
        test_pdf.set_font('Arial', 'B', 16)
        test_pdf.cell(0, 10, f"Тест - Варіант {variant['variant_number']}", ln=True, align='C')
        test_pdf.ln(10)
        
        test_pdf.set_font('Arial', '', 12)
        test_pdf.cell(0, 8, "Інструкція: Оберіть правильну відповідь і впишіть її номер у таблицю внизу.", ln=True)
        test_pdf.ln(5)
        
        # Додаємо питання з підтримкою колонок
        columns = max(1, min(3, columns))  # Обмежуємо від 1 до 3 колонок
        questions = variant['questions']
        
        if columns == 1:
            # Одна колонка - просте розміщення
            for q_idx, question in enumerate(questions):
                q_num = q_idx + 1
                
                # Перевіряємо місце на сторінці
                if test_pdf.get_y() > 220:  # Залишаємо місце для таблиці відповідей
                    test_pdf.add_page()
                    test_pdf.set_font('Arial', 'B', 16)
                    test_pdf.cell(0, 10, f"Тест - Варіант {variant['variant_number']} (продовження)", ln=True, align='C')
                    test_pdf.ln(10)
                
                test_pdf.set_font('Arial', 'B', 11)
                question_text = f"{q_num}. {question['question_text']}"
                add_multiline_text(test_pdf, question_text, page_width, 6, 9)
                
                test_pdf.set_font('Arial', '', 10)
                for opt_num, option in enumerate(question['options'], 1):
                    option_text = f"   {opt_num}) {option}"
                    add_multiline_text(test_pdf, option_text, page_width, 5, 8)
                
                test_pdf.ln(3)
        else:
            # Багатоколонкове розміщення
            column_width = (page_width - (columns - 1) * 10) / columns
            
            # Масив для відстеження Y позицій кожної колонки
            column_y_positions = [test_pdf.get_y()] * columns
            start_y = test_pdf.get_y()
            
            for q_idx, question in enumerate(questions):
                q_num = q_idx + 1
                col = q_idx % columns
                x_pos = test_pdf.l_margin + col * (column_width + 10)
                
                # Перевіряємо місце на сторінці - використовуємо максимальну Y позицію серед усіх колонок
                max_y = max(column_y_positions)
                if max_y > 200:  # Зменшуємо поріг для кращого контролю
                    test_pdf.add_page()
                    test_pdf.set_font('Arial', 'B', 16)
                    test_pdf.cell(0, 10, f"Тест - Варіант {variant['variant_number']} (продовження)", ln=True, align='C')
                    test_pdf.ln(10)
                    column_y_positions = [test_pdf.get_y()] * columns
                    start_y = test_pdf.get_y()
                
                # Встановлюємо позицію для поточної колонки
                test_pdf.set_xy(x_pos, column_y_positions[col])
                
                test_pdf.set_font('Arial', 'B', 10)
                question_text = f"{q_num}. {question['question_text']}"
                
                # Розбиваємо текст на рядки
                fitted_text, font_size = fit_text_to_width(test_pdf, question_text, column_width, 8)
                test_pdf.set_font('Arial', 'B', font_size)
                
                lines = fitted_text.split('\n')
                for line in lines:
                    test_pdf.cell(column_width, 5, line, ln=True)
                    test_pdf.set_x(x_pos)
                
                test_pdf.set_font('Arial', '', 9)
                for opt_num, option in enumerate(question['options'], 1):
                    option_text = f"   {opt_num}) {option}"
                    fitted_option, opt_font_size = fit_text_to_width(test_pdf, option_text, column_width, 7)
                    test_pdf.set_font('Arial', '', opt_font_size)
                    
                    opt_lines = fitted_option.split('\n')
                    for opt_line in opt_lines:
                        test_pdf.cell(column_width, 4, opt_line, ln=True)
                        test_pdf.set_x(x_pos)
                
                # Оновлюємо Y позицію для поточної колонки
                column_y_positions[col] = test_pdf.get_y() + 3
            
            # Встановлюємо позицію після всіх колонок
            test_pdf.set_xy(test_pdf.l_margin, max(column_y_positions))
        
        # Додаємо таблицю для відповідей
        test_pdf.ln(8)
        test_pdf.set_font('Arial', 'B', 11)
        test_pdf.cell(0, 6, "Таблиця відповідей:", ln=True)
        test_pdf.ln(3)
        
        # Створюємо компактну таблицю відповідей з однаковими ячейками
        test_pdf.set_font('Arial', '', 9)
        # Встановлюємо сірий колір для рамок D3D3D3 (RGB: 211, 211, 211)
        test_pdf.set_draw_color(211, 211, 211)
        questions_per_row = 15  # Кількість питань у рядку
        num_questions = len(variant['questions'])
        
        # Фіксована ширина ячейки для рівномірності
        fixed_cell_width = page_width / questions_per_row
        
        for row_start in range(0, num_questions, questions_per_row):
            questions_in_row = min(questions_per_row, num_questions - row_start)
            
            # Номери питань - всі ячейки однакового розміру
            test_pdf.set_font('Arial', 'B', 8)
            for i in range(questions_in_row):
                test_pdf.cell(fixed_cell_width, 6, f"№{row_start + i + 1}", 1, 0, 'C')
            # Заповнюємо решту ячейок порожніми для вирівнювання
            for i in range(questions_in_row, questions_per_row):
                test_pdf.cell(fixed_cell_width, 6, "", 1, 0, 'C')
            test_pdf.ln()
            
            # Порожні ячейки для відповідей - всі однакового розміру
            test_pdf.set_font('Arial', '', 8)
            for i in range(questions_in_row):
                test_pdf.cell(fixed_cell_width, 8, "", 1, 0, 'C')
            # Заповнюємо решту ячейок порожніми для вирівнювання
            for i in range(questions_in_row, questions_per_row):
                test_pdf.cell(fixed_cell_width, 8, "", 1, 0, 'C')
            test_pdf.ln()
            test_pdf.ln(2)
    
    test_pdf.output(test_pdf_path)
    
    # Створюємо PDF з відповідями для вчителя
    answers_pdf = FPDF()
    answers_pdf.add_font('Arial', '', 'c:/windows/fonts/arial.ttf', uni=True)
    answers_pdf.add_font('Arial', 'B', 'c:/windows/fonts/arialbd.ttf', uni=True)
    
    # Встановлюємо сірий колір для рамок D3D3D3 (RGB: 211, 211, 211)
    answers_pdf.set_draw_color(211, 211, 211)
    
    answer_page_width = answers_pdf.w - 2 * answers_pdf.l_margin
    
    answers_pdf.add_page()
    answers_pdf.set_font('Arial', 'B', 16)
    answers_pdf.cell(0, 10, "Відповіді для вчителя", ln=True, align='C')
    answers_pdf.ln(10)
    
    # Групуємо варіанти по більше на сторінку
    variants_per_page = 8  # Збільшуємо кількість варіантів на сторінці
    for page_start in range(0, len(variants), variants_per_page):
        if page_start > 0:
            answers_pdf.add_page()
        
        for variant in variants[page_start:page_start + variants_per_page]:
            answers_pdf.set_font('Arial', 'B', 12)
            variant_text = f"Варіант {variant['variant_number']}"
            add_multiline_text(answers_pdf, variant_text, answer_page_width, 6, 9)
            
            answers_pdf.set_font('Arial', '', 10)
            # Виводимо відповіді більш компактно
            answer_text = "Відповіді: " + ", ".join([f"{i+1}-{ans}" for i, ans in enumerate(variant['answer_key'])])
            add_multiline_text(answers_pdf, answer_text, answer_page_width, 5, 8)
            answers_pdf.ln(3)  # Зменшуємо відступ між варіантами
    
    answers_pdf.output(answers_pdf_path)
    
    log.info(f"Створено PDF файли: {test_pdf_path}, {answers_pdf_path}")
    return test_pdf_path, answers_pdf_path

def create_excel_answer_key(variants: List[Dict[str, Any]], output_dir: str, input_file_name: str = "") -> str:
    """
    Створює Excel файл-ключ з відповідями для всіх варіантів.
    
    Args:
        variants: Список варіантів тестів
        output_dir: Папка для збереження файлу
        
    Returns:
        Шлях до створеного Excel файлу
    """
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    if input_file_name:
        excel_path = os.path.join(output_dir, f"{input_file_name}_ключ_{timestamp}.xlsx")
    else:
        excel_path = os.path.join(output_dir, f"answer_key_{timestamp}.xlsx")
    
    # Створюємо Excel файл з кількома аркушами
    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        # Аркуш 1: Основні відповіді (для зворотної сумісності)
        basic_data = []
        for variant in variants:
            # Створюємо рядок з відповідями через кому
            answers_str = ",".join(map(str, variant['answer_key']))
            # Створюємо рядок з вагами через кому
            weights_str = ",".join(str(q['weight']) for q in variant['questions'])
            basic_data.append({
                'Варіант': variant['variant_number'],
                'Відповіді': answers_str,
                'Ваги': weights_str
            })
        
        basic_df = pd.DataFrame(basic_data)
        basic_df.to_excel(writer, sheet_name='Основні_відповіді', index=False)
        
        # Аркуш 2: Детальна інформація про питання
        detailed_data = []
        for variant in variants:
            for i, question in enumerate(variant['questions']):
                question_data = {
                    'Варіант': variant['variant_number'],
                    'Номер_питання': i + 1,
                    'Текст_питання': question['question_text'],
                    'Тип_питання': 'Тестове' if question['is_test_question'] else 'Відкрите',
                    'Правильна_відповідь': question['correct_answer'],
                    'Вага': question['weight']
                }
                
                # Додаємо варіанти відповідей для тестових питань
                if question['is_test_question'] and 'options' in question:
                    for j, option in enumerate(question['options']):
                        question_data[f'Варіант_{j+1}'] = option
                
                detailed_data.append(question_data)
        
        detailed_df = pd.DataFrame(detailed_data)
        detailed_df.to_excel(writer, sheet_name='Детальна_інформація', index=False)
        
        # Аркуш 3: Статистика варіантів
        stats_data = []
        for variant in variants:
            total_questions = len(variant['questions'])
            test_questions = sum(1 for q in variant['questions'] if q['is_test_question'])
            open_questions = total_questions - test_questions
            total_weight = sum(q['weight'] for q in variant['questions'])
            avg_weight = total_weight / total_questions if total_questions > 0 else 0
            
            stats_data.append({
                'Варіант': variant['variant_number'],
                'Загальна_кількість_питань': total_questions,
                'Тестових_питань': test_questions,
                'Відкритих_питань': open_questions,
                'Загальна_вага': total_weight,
                'Середня_вага': round(avg_weight, 2)
            })
        
        stats_df = pd.DataFrame(stats_data)
        stats_df.to_excel(writer, sheet_name='Статистика_варіантів', index=False)
    
    log.info(f"Створено розширений Excel файл-ключ: {excel_path}")
    return excel_path

def check_student_answers(answer_key_file: str, variant_number: int, student_answers: List) -> Dict[str, Any]:
    """
    Перевіряє відповіді учня за файлом-ключем.
    
    Args:
        answer_key_file: Шлях до Excel файлу-ключа
        variant_number: Номер варіанту учня
        student_answers: Список відповідей учня (може містити числа та рядки)
        
    Returns:
        Словник з результатами перевірки
    """
    try:
        # Спробуємо прочитати розширений формат (з кількома аркушами)
        detailed_info = None
        try:
            # Читаємо основний аркуш
            key_df = pd.read_excel(answer_key_file, sheet_name='Основні_відповіді')
            # Читаємо детальну інформацію
            detailed_df = pd.read_excel(answer_key_file, sheet_name='Детальна_інформація')
            detailed_info = detailed_df[detailed_df['Варіант'] == variant_number]
        except:
            # Якщо не вдалося прочитати розширений формат, читаємо старий формат
            key_df = pd.read_excel(answer_key_file)
        
        # Знаходимо рядок з потрібним варіантом
        variant_row = key_df[key_df['Варіант'] == variant_number]
        if variant_row.empty:
            raise ValueError(f"Варіант {variant_number} не знайдено у файлі-ключі")
        
        # Витягуємо відповіді та ваги
        answers_str = variant_row['Відповіді'].iloc[0]
        weights_str = variant_row['Ваги'].iloc[0]
        
        # Парсимо відповіді та ваги
        answer_key = []
        weights = []
        
        for ans in str(answers_str).split(','):
            ans = ans.strip()
            # Зберігаємо відповіді як рядки, щоб зберегти початковий формат
            # Це дозволить розрізняти '3' і '03' для відкритих завдань
            answer_key.append(ans)
        
        for weight in str(weights_str).split(','):
            weights.append(float(weight.strip()))
        
        # Перевіряємо кількість відповідей
        if len(student_answers) != len(answer_key):
            raise ValueError(f"Кількість відповідей учня ({len(student_answers)}) не збігається з кількістю питань ({len(answer_key)})")
        
        # Підраховуємо правильні відповіді з урахуванням ваг
        total_weight = sum(weights)
        total_points = 12  # Загальна кількість балів за тест
        correct_weighted_score = 0
        detailed_results = []
        
        for i, (student_ans, correct_ans, weight) in enumerate(zip(student_answers, answer_key, weights)):
            question_points = (weight / total_weight) * total_points
            
            # Отримуємо додаткову інформацію про питання, якщо доступна
            question_text = ""
            question_type = ""
            question_options = []
            
            if detailed_info is not None and not detailed_info.empty:
                question_detail = detailed_info[detailed_info['Номер_питання'] == i + 1]
                if not question_detail.empty:
                    question_text = question_detail['Текст_питання'].iloc[0]
                    question_type = question_detail['Тип_питання'].iloc[0]
                    # Збираємо варіанти відповідей для тестових питань
                    for col in question_detail.columns:
                        if col.startswith('Варіант_') and pd.notna(question_detail[col].iloc[0]):
                            question_options.append(str(question_detail[col].iloc[0]))
            
            # Перевіряємо правильність відповіді
            # Визначаємо тип питання: тестове питання має варіанти відповідей, відкрите - не має
            is_test_question = len(question_options) > 0
            
            if is_test_question:
                # Тестове завдання - порівнюємо літери
                # Перевіряємо чи відповідь не порожня
                if not student_ans or str(student_ans).strip() == "":
                    # Для тестових питань порожня відповідь вважається неправильною
                    is_correct = False
                    student_ans_int = "(не заповнено)"
                else:
                    # Українські літери для варіантів відповідей (без Ґ, Є, І, Ї, Й, Ь)
                    ukrainian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ю', 'Я']
                    
                    # Нормалізуємо відповіді учня та правильну відповідь
                    student_ans_normalized = str(student_ans).strip().upper()
                    correct_ans_normalized = str(correct_ans).strip().upper()
                    
                    num_options = len(question_options)
                    # Використовуємо всі доступні українські літери для валідації
                    valid_letters = ukrainian_letters
                    
                    # Перевіряємо множинний вибір (кілька літер)
                    if len(student_ans_normalized) > 1:
                        # Множинний вибір - перевіряємо всі літери
                        student_letters = list(student_ans_normalized)
                        correct_letters = list(correct_ans_normalized)
                        
                        # Перевіряємо, що всі літери учня є валідними
                        invalid_letters = [letter for letter in student_letters if letter not in valid_letters]
                        if invalid_letters:
                            raise ValueError(f"Питання {i+1}: літери '{', '.join(invalid_letters)}' не підходять. Доступні літери: {', '.join(valid_letters)}")
                        
                        # Множинний вибір - строга перевірка точної відповідності
                        student_set = set(student_letters)
                        correct_set = set(correct_letters)
                        
                        # Бінарне оцінювання: або повністю правильно, або 0 балів
                        is_correct = student_set == correct_set
                        student_ans_int = student_ans_normalized
                    else:
                        # Одиночний вибір
                        if student_ans_normalized not in valid_letters:
                            raise ValueError(f"Питання {i+1}: літера '{student_ans}' не підходить. Доступні літери: {', '.join(valid_letters)}")
                        
                        is_correct = student_ans_normalized == correct_ans_normalized
                        student_ans_int = student_ans_normalized
            else:
                # Відкрите завдання - порівнюємо рядки з нормалізацією
                # Перевіряємо чи відповідь не порожня
                if not student_ans or str(student_ans).strip() == "":
                    # Для відкритих питань порожня відповідь вважається неправильною
                    is_correct = False
                    student_ans_int = "(не заповнено)"
                else:
                    student_str = str(student_ans).strip().lower()
                    correct_str = str(correct_ans).strip().lower()
                    is_correct = student_str == correct_str
                    student_ans_int = student_ans
            
            # Підраховуємо бали - бінарне оцінювання
            earned_points = question_points if is_correct else 0
            
            correct_weighted_score += earned_points
            
            result_item = {
                'question_number': i + 1,
                'student_answer': student_ans_int,
                'correct_answer': correct_ans,
                'is_correct': is_correct,
                'weight': weight,
                'points': earned_points,
                'max_points': question_points,
                'is_test_question': is_test_question,
                'question_text': question_text,
                'question_type': question_type,
                'question_options': question_options
            }
            
            detailed_results.append(result_item)
        
        # Обчислюємо відсоток
        score_percentage = (correct_weighted_score / total_points) * 100
        correct_count = sum(1 for r in detailed_results if r['is_correct'])
        
        result = {
            'variant_number': variant_number,
            'total_questions': len(answer_key),
            'correct_answers': correct_count,
            'score_percentage': score_percentage,
            'weighted_score': correct_weighted_score,
            'max_score': total_points,
            'detailed_results': detailed_results
        }
        
        log.info(f"Перевірка завершена для варіанта {variant_number}: {correct_count}/{len(answer_key)} ({format_number_with_comma(score_percentage, 1)}%, {format_number_with_comma(correct_weighted_score, 2)}/{total_points} балів)")
        return result
        
    except Exception as e:
        log.error(f"Помилка при перевірці відповідей: {e}")
        raise

def create_check_result_pdf(check_result: Dict[str, Any], output_dir: str) -> str:
    """
    Створює PDF файл з результатами перевірки.
    
    Args:
        check_result: Результати перевірки
        output_dir: Папка для збереження файлу
        
    Returns:
        Шлях до створеного PDF файлу
    """
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    
    # Нормалізуємо шлях і створюємо папку
    try:
        output_dir = os.path.normpath(output_dir)
        os.makedirs(output_dir, exist_ok=True)
    except Exception as e:
        log.warning(f"Помилка при роботі з папкою {output_dir}: {e}. Використовуємо тимчасову папку.")
        output_dir = ensure_temp_dir("reports_")
    
    # Формуємо ім'я файлу у форматі Клас_ПІБ_Варіант_Дата
    student_info = check_result.get('student_info', {})
    class_name = student_info.get('class', '').replace(' ', '_').replace('-', '_') or 'БезКласу'
    full_name = student_info.get('full_name', '').replace(' ', '_') or 'БезІмені'
    variant = check_result['variant_number']
    
    filename = f"{class_name}_{full_name}_Варіант{variant}_{timestamp}.pdf"
    pdf_path = os.path.join(output_dir, filename)
    
    pdf = FPDF()
    pdf.add_font('Arial', '', 'c:/windows/fonts/arial.ttf', uni=True)
    pdf.add_font('Arial', 'B', 'c:/windows/fonts/arialbd.ttf', uni=True)
    
    check_page_width = pdf.w - 2 * pdf.l_margin
    
    pdf.add_page()
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, "Результат перевірки тесту (короткий звіт)", ln=True, align='C')
    pdf.ln(10)
    
    # Дані учня (якщо є)
    student_info = check_result.get('student_info', {})
    if any(student_info.values()):
        pdf.set_font('Arial', 'B', 12)
        add_multiline_text(pdf, "Дані учня / учениці:", check_page_width, 8, 10)
        pdf.set_font('Arial', '', 10)
        
        if student_info.get('class'):
            add_multiline_text(pdf, f"Клас: {student_info['class']}", check_page_width, 6, 10)
        if student_info.get('full_name'):
            add_multiline_text(pdf, f"ПІБ: {student_info['full_name']}", check_page_width, 6, 10)
        pdf.ln(5)
    
    # Основна інформація
    pdf.set_font('Arial', 'B', 12)
    # Використовуємо зважені бали
    weighted_score = check_result.get('weighted_score', 0)
    max_score = check_result.get('max_score', 12)
    info_texts = [
        f"Варіант: {check_result['variant_number']}",
        f"Всього питань: {check_result['total_questions']}",
        f"Правильних відповідей: {check_result['correct_answers']}",
        f"Відсоток: {format_number_with_comma(check_result['score_percentage'], 1)}%",
        f"Бали: {format_number_with_comma(weighted_score, 2)} з {max_score}"
    ]
    
    for info_text in info_texts:
        add_multiline_text(pdf, info_text, check_page_width, 8, 10)
    pdf.ln(10)
    
    # Детальные результаты
    pdf.set_font('Arial', 'B', 12)
    add_multiline_text(pdf, "Детальні результати:", check_page_width, 8, 10)
    pdf.ln(5)
    
    # Заголовки таблицы
    pdf.set_font('Arial', 'B', 10)
    pdf.cell(20, 8, "Питання", 1, 0, 'C')
    pdf.cell(30, 8, "Відповідь учня", 1, 0, 'C')
    pdf.cell(45, 8, "Правильна відповідь", 1, 0, 'C')
    pdf.cell(30, 8, "Бали", 1, 0, 'C')
    pdf.cell(35, 8, "Результат", 1, 0, 'C')
    pdf.ln()
    
    # Строки таблицы
    pdf.set_font('Arial', '', 10)
    for result in check_result['detailed_results']:
        pdf.cell(20, 8, str(result['question_number']), 1, 0, 'C')
        pdf.cell(30, 8, str(result['student_answer']), 1, 0, 'C')
        pdf.cell(45, 8, str(result['correct_answer']), 1, 0, 'C')
        # Баллы за задание с учетом весов
        earned_points = result.get('points', 0)
        max_points = result.get('max_points', 0)
        points_text = f"{format_number_with_comma(earned_points, 2)}/{format_number_with_comma(max_points, 2)}"
        pdf.cell(30, 8, points_text, 1, 0, 'C')
        # Используем текст вместо символов, которые не поддерживаются шрифтом Arial
        result_text = "Правильно" if result['is_correct'] else "Неправильно"
        pdf.cell(35, 8, result_text, 1, 0, 'C')
        pdf.ln()
    
    try:
        pdf.output(pdf_path)
        log.info(f"Створено PDF з результатами перевірки: {pdf_path}")
        return pdf_path
    except Exception as e:
        # Если ошибка связана с путем, пробуем использовать временную папку
        if 'Invalid argument' in str(e) or 'path' in str(e).lower():
            try:
                output_dir = ensure_temp_dir("reports_")
                log.warning(f"Помилка зі шляхом PDF, використовуємо тимчасову папку: {output_dir}")
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                pdf_path = os.path.join(output_dir, f"check_result_variant_{check_result['variant_number']}_{timestamp}.pdf")
                pdf.output(pdf_path)
                log.info(f"Створено PDF з результатами перевірки (fallback): {pdf_path}")
                return pdf_path
            except Exception as fallback_error:
                log.error(f"Помилка при створенні PDF звіту (fallback): {fallback_error}")
                raise
        else:
            log.error(f"Помилка при створенні PDF звіту: {e}")
            raise

def create_check_result_word(check_result: Dict[str, Any], output_dir: str) -> str:
    """
    Создает Word файл с результатами проверки.
    
    Args:
        check_result: Результаты проверки
        output_dir: Папка для сохранения файла
        
    Returns:
        Путь к созданному Word файлу
    """
    try:
        # Нормализуем путь и создаем папку
        output_dir = os.path.normpath(output_dir)
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        
        # Формируем имя файла в формате Класс_ПІБ_Вариант_Дата
        student_info = check_result.get('student_info', {})
        class_name = student_info.get('class', '').replace(' ', '_').replace('-', '_') or 'БезКласу'
        full_name = student_info.get('full_name', '').replace(' ', '_') or 'БезІмені'
        variant = check_result['variant_number']
        
        filename = f"{class_name}_{full_name}_Варіант{variant}_{timestamp}.docx"
        word_path = os.path.join(output_dir, filename)
        
        doc = Document()
        
        # Устанавливаем минимальные поля для экономии места
        section = doc.sections[0]
        section.top_margin = Inches(0.5)     # 1.27 см
        section.bottom_margin = Inches(0.5)  # 1.27 см
        section.left_margin = Inches(0.5)    # 1.27 см
        section.right_margin = Inches(0.5)   # 1.27 см
        
        # Заголовок
        heading = doc.add_heading('Результат перевірки тесту (розширений звіт)', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Пустая строка
        
        # Данные ученика (если есть)
        student_info = check_result.get('student_info', {})
        if any(student_info.values()):
            student_para = doc.add_paragraph()
            student_para.add_run('Дані учня / учениці:').bold = True
            
            if student_info.get('class'):
                doc.add_paragraph(f"Клас: {student_info['class']}", style='List Bullet')
            if student_info.get('full_name'):
                doc.add_paragraph(f"ПІБ: {student_info['full_name']}", style='List Bullet')
            
            doc.add_paragraph()  # Пустая строка
        
        # Основная информация
        info_para = doc.add_paragraph()
        info_para.add_run('Основна інформація:').bold = True
        
        # Используем взвешенные баллы
        weighted_score = check_result.get('weighted_score', 0)
        max_score = check_result.get('max_score', 12)
        info_texts = [
            f"Варіант: {check_result['variant_number']}",
            f"Всього питань: {check_result['total_questions']}",
            f"Правильних відповідей: {check_result['correct_answers']}",
            f"Відсоток: {format_number_with_comma(check_result['score_percentage'], 1)}%",
        f"Бали: {format_number_with_comma(weighted_score, 2)} з {max_score}"
        ]
        
        for info_text in info_texts:
            doc.add_paragraph(info_text, style='List Bullet')
        
        doc.add_paragraph()  # Пустая строка
        
        # Детальные результаты
        details_para = doc.add_paragraph()
        details_para.add_run('Детальні результати:').bold = True
        
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Настраиваем ширину колонок
        table.columns[0].width = Inches(1.0)  # Питання
        table.columns[1].width = Inches(1.5)  # Відповідь учня
        table.columns[2].width = Inches(1.5)  # Правильна відповідь
        table.columns[3].width = Inches(0.8)  # Бали
        table.columns[4].width = Inches(1.8)  # Результат (увеличена для предотвращения переносов)
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Питання'
        hdr_cells[1].text = 'Відповідь учня'
        hdr_cells[2].text = 'Правильна відповідь'
        hdr_cells[3].text = 'Бали'
        hdr_cells[4].text = 'Результат'
        
        # Делаем заголовки жирными
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Добавляем строки с результатами
        for result in check_result['detailed_results']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(result['question_number'])
            row_cells[1].text = str(result['student_answer'])
            row_cells[2].text = str(result['correct_answer'])
            # Баллы за задание с учетом веса
            earned_points = result.get('points', 0)
            max_points = result.get('max_points', 0)
            points = f"{format_number_with_comma(earned_points, 2)} / {format_number_with_comma(max_points, 2)}"
            row_cells[3].text = points
            
            # Результат с цветными символами
            result_paragraph = row_cells[4].paragraphs[0]
            result_paragraph.clear()
            if result['is_correct']:
                run = result_paragraph.add_run("✓ Правильно")
                run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый цвет
            else:
                run = result_paragraph.add_run("✗ Неправильно")
                run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет
        
        # Добавляем серые рамки D3D3D3 ко всем ячейкам таблицы
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/></w:tcBorders>'))
        
        # Добавляем детальный анализ с полным текстом вопросов
        doc.add_page_break()
        
        # Заголовок детального анализа
        detailed_heading = doc.add_heading('Детальний аналіз відповідей', level=2)
        detailed_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Пустая строка
        
        # Проходим по каждому вопросу
        for i, result in enumerate(check_result['detailed_results'], 1):
            # Заголовок вопроса
            question_heading = doc.add_heading(f'Питання {result["question_number"]}', level=3)
            
            # Текст вопроса
            question_text = result.get('question_text', f'Питання {i+1}')
            if question_text:
                question_para = doc.add_paragraph()
                question_para.add_run('Текст питання: ').bold = True
                add_formatted_text_to_paragraph(question_para, question_text)
            

            
            # Тип вопроса з індикатором
            question_type = result.get('question_type', 'Невідомий')
            if not question_type and result.get('is_test_question'):
                question_type = 'Тестове'
            elif not question_type:
                question_type = 'Відкрите'
            
            # Створюємо об'єкт питання для get_task_type_indicator
            question_obj = {
                'is_test_question': result.get('is_test_question', True),
                'correct_answer': result.get('correct_answer', '')
            }
            task_indicator = get_task_type_indicator(question_obj)
            
            type_para = doc.add_paragraph()
            type_para.add_run('Тип питання: ').bold = True
            type_para.add_run(f"{question_type} {task_indicator}")
            
            # Варианты ответов для тестовых вопросов в виде трьохколонкової таблиці
            question_options = result.get('question_options', [])
            if question_options and result.get('is_test_question'):
                options_para = doc.add_paragraph()
                options_para.add_run('Варіанти відповідей:').bold = True
                
                # Створюємо таблицю без рамок з трьома колонками
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Додаємо світло-сірі границі таблиці
                for row in table.rows:
                     for cell in row.cells:
                         cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/></w:tcBorders>'))
                
                # Для тестових питань відповіді можуть бути українськими літерами
                student_answer_str = str(result['student_answer']).strip().upper()
                correct_answer_str = str(result['correct_answer']).strip().upper()
                
                # Перевіряємо чи відповідь порожня (не заповнено)
                if student_answer_str == "(НЕ ЗАПОВНЕНО)":
                    student_answer_str = ""
                
                for j, option in enumerate(question_options, 1):
                    # Форматируем число правильно, используя ту же логику что и при генерации
                    def format_option_value(value):
                        if isinstance(value, (int, float)):
                            try:
                                # Проверяем, является ли число целым
                                if float(value) == int(float(value)):
                                    return str(int(float(value)))
                                else:
                                    return f"{float(value):.10g}".replace('.', ',')
                            except (ValueError, TypeError):
                                return str(value).strip()
                        else:
                            # Для строковых значений проверяем, является ли это украинской буквой
                            str_value = str(value).strip()
                            ukrainian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ю', 'Я']
                            if str_value.upper() in ukrainian_letters:
                                return str_value.upper()
                            # Пытаемся преобразовать в число только если это не украинская буква
                            try:
                                num_value = float(str_value.replace(',', '.'))
                                if num_value == int(num_value):
                                    return str(int(num_value))
                                else:
                                    return f"{num_value:.10g}".replace('.', ',')
                            except (ValueError, TypeError):
                                return str_value
                    
                    formatted_option = format_option_value(option)
                    
                    # Додаємо рядок до таблиці
                    row_cells = table.add_row().cells
                    
                    # Перетворюємо номер варіанту на українську літеру або номер
                    ukrainian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ю', 'Я']
                    option_letter = ukrainian_letters[j-1] if j <= len(ukrainian_letters) else str(j)
                    
                    # Перша колонка: номер завдання і варіант
                    row_cells[0].text = f'{option_letter}. {formatted_option}'
                    
                    # Перевіряємо чи цей варіант обрав учень або чи це правильна відповідь
                    is_student_choice = False
                    is_correct_choice = False
                    
                    # Для множинного вибору перевіряємо кожну літеру окремо
                    if len(student_answer_str) > 1:  # Множинний вибір
                        is_student_choice = option_letter in student_answer_str
                    else:  # Одиночний вибір
                        is_student_choice = option_letter == student_answer_str
                    
                    if len(correct_answer_str) > 1:  # Множинний вибір
                        is_correct_choice = option_letter in correct_answer_str
                    else:  # Одиночний вибір
                        is_correct_choice = option_letter == correct_answer_str
                    
                    # Друга колонка: вибір учня (без кольорових виділень)
                    if is_student_choice:
                        row_cells[1].text = '✓ Обрано'
                    else:
                        row_cells[1].text = ''
                    
                    # Третя колонка: позначка правильних і неправильних відповідей (з кольором)
                    if is_student_choice and is_correct_choice:
                        # Правильно обрано
                        paragraph = row_cells[2].paragraphs[0]
                        run = paragraph.add_run('✓ ПРАВИЛЬНО')
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Зелений
                        run.bold = True
                    elif is_student_choice and not is_correct_choice:
                        # Неправильно обрано
                        paragraph = row_cells[2].paragraphs[0]
                        run = paragraph.add_run('✗ НЕПРАВИЛЬНО')
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Червоний
                        run.bold = True
                    elif not is_student_choice and is_correct_choice:
                        # Правильна відповідь, але не обрана
                        paragraph = row_cells[2].paragraphs[0]
                        run = paragraph.add_run('✓ Правильна відповідь')
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Зелений
                        run.bold = True
                    else:
                        row_cells[2].text = ''
                    
                    # Додаємо світло-сірі границі для нового рядка
                    for cell in row_cells:
                        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/></w:tcBorders>'))
                
                # Видаляємо перший порожній рядок
                table._element.remove(table.rows[0]._element)
            
            # Додаємо відображення відповідей для тестових питань
            if result.get('is_test_question'):
                answers_para = doc.add_paragraph()
                
                # Форматуємо відповіді учня та правильну відповідь
                student_answer_display = str(result['student_answer']).strip()
                correct_answer_display = str(result['correct_answer']).strip()
                
                # Перевіряємо чи відповідь порожня
                if student_answer_display == "(НЕ ЗАПОВНЕНО)" or student_answer_display == "":
                    student_answer_display = "(не заповнено)"
                
                answers_para.add_run('Відповіді: ').bold = True
                answers_para.add_run('Учень відповів: ')
                
                student_run = answers_para.add_run(f'({student_answer_display})')
                if result['is_correct']:
                    student_run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
                else:
                    student_run.font.color.rgb = RGBColor(255, 0, 0)  # Красный
                student_run.bold = True
                
                answers_para.add_run(', правильна відповідь: ')
                correct_run = answers_para.add_run(f'({correct_answer_display})')
                correct_run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
                correct_run.bold = True
            
            # Для відкритих питань додаємо інформацію про відповіді
            if not result.get('is_test_question'):
                answers_para = doc.add_paragraph()
                
                # Форматируем ответы правильно, используя ту же логику что и при генерации
                student_answer = result['student_answer']
                correct_answer = result['correct_answer']
                
                # Применяем правильное форматирование чисел как при генерации вариантов
                def format_answer_value(value):
                    if isinstance(value, (int, float)):
                        try:
                            # Проверяем, является ли число целым
                            if float(value) == int(float(value)):
                                return str(int(float(value)))
                            else:
                                return f"{float(value):.10g}".replace('.', ',')
                        except (ValueError, TypeError):
                            return str(value).strip()
                    else:
                        # Для строковых значений проверяем, является ли это украинской буквой
                        str_value = str(value).strip()
                        ukrainian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ю', 'Я']
                        if str_value.upper() in ukrainian_letters:
                            return str_value.upper()
                        # Пытаемся преобразовать в число только если это не украинская буква
                        try:
                            num_value = float(str_value.replace(',', '.'))
                            if num_value == int(num_value):
                                return str(int(num_value))
                            else:
                                return f"{num_value:.10g}".replace('.', ',')
                        except (ValueError, TypeError):
                            return str_value
                
                formatted_student = format_answer_value(student_answer)
                formatted_correct = format_answer_value(correct_answer)
                
                answers_para.add_run('Відповіді: ').bold = True
                answers_para.add_run('Учень відповів: ')
                
                student_run = answers_para.add_run(f'({formatted_student})')
                if result['is_correct']:
                    student_run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
                else:
                    student_run.font.color.rgb = RGBColor(255, 0, 0)  # Красный
                student_run.bold = True
                
                answers_para.add_run(', правильна відповідь: ')
                correct_run = answers_para.add_run(f'({formatted_correct})')
                correct_run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
                correct_run.bold = True
            
            # Результат с цветовым выделением
            result_para = doc.add_paragraph()
            result_para.add_run('Результат: ').bold = True
            
            if result['is_correct']:
                 result_run = result_para.add_run('✓ ПРАВИЛЬНО')
                 result_run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
                 result_run.bold = True
            else:
                result_run = result_para.add_run('✗ НЕПРАВИЛЬНО')
                result_run.font.color.rgb = RGBColor(255, 0, 0)  # Красный
                result_run.bold = True
            
            # Баллы
            earned_points = result.get('points', 0)
            max_points = result.get('max_points', 0)
            points_para = doc.add_paragraph()
            points_para.add_run('Бали: ').bold = True
            points_text = f'{format_number_with_comma(earned_points, 2)} з {format_number_with_comma(max_points, 2)}'
            points_run = points_para.add_run(points_text)
            
            if result['is_correct']:
                points_run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
            else:
                points_run.font.color.rgb = RGBColor(255, 0, 0)  # Красный
            points_run.bold = True
            

            
            # Разделитель между вопросами
            if i < len(check_result['detailed_results']):
                doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Пустая строка
        
        doc.save(word_path)
        log.info(f"Створено Word документ з результатами перевірки: {word_path}")
        return word_path
        
    except Exception as e:
        # Если ошибка связана с путем, пробуем использовать временную папку
        if 'Invalid argument' in str(e) or 'path' in str(e).lower():
            try:
                output_dir = ensure_temp_dir("reports_")
                log.warning(f"Помилка зі шляхом, використовуємо тимчасову папку: {output_dir}")
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                word_path = os.path.join(output_dir, f"check_result_variant_{check_result['variant_number']}_{timestamp}.docx")
                
                doc = Document()
                
                # Заголовок
                heading = doc.add_heading('Результат перевірки тесту', level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()  # Пустая строка
                
                # Основная информация
                info_para = doc.add_paragraph()
                info_para.add_run('Основна інформація:').bold = True
                
                # Расчет балла в 12-балльной системе
                score_12 = (check_result['correct_answers'] / check_result['total_questions']) * 12
                
                info_texts = [
                    f"Варіант: {check_result['variant_number']}",
                    f"Всього питань: {check_result['total_questions']}",
                    f"Правильних відповідей: {check_result['correct_answers']}",
                    f"Відсоток: {format_number_with_comma(check_result['score_percentage'], 1)}%",
            f"Оцінка (12-бальна система): {format_number_with_comma(score_12, 2)}"
                ]
                
                for info_text in info_texts:
                    doc.add_paragraph(info_text, style='List Bullet')
                
                doc.add_paragraph()  # Пустая строка
                
                # Детальные результаты
                details_para = doc.add_paragraph()
                details_para.add_run('Детальні результати:').bold = True
                
                # Создаем таблицу
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Настраиваем ширину колонок
                table.columns[0].width = Inches(1.0)  # Питання
                table.columns[1].width = Inches(1.5)  # Відповідь учня
                table.columns[2].width = Inches(1.5)  # Правильна відповідь
                table.columns[3].width = Inches(0.8)  # Бали
                table.columns[4].width = Inches(1.8)  # Результат (увеличена для предотвращения переносов)
                
                # Заголовки таблицы
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Питання'
                hdr_cells[1].text = 'Відповідь учня'
                hdr_cells[2].text = 'Правильна відповідь'
                hdr_cells[3].text = 'Бали'
                hdr_cells[4].text = 'Результат'
                
                # Делаем заголовки жирными
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Добавляем строки с результатами
                for result in check_result['detailed_results']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(result['question_number'])
                    row_cells[1].text = str(result['student_answer'])
                    row_cells[2].text = str(result['correct_answer'])
                    # Баллы за задание с учетом веса
                    earned_points = result.get('points', 0)
                    max_points = result.get('max_points', 0)
                    points = f"{format_number_with_comma(earned_points, 2)} / {format_number_with_comma(max_points, 2)}"
                    row_cells[3].text = points
                    
                    # Результат с цветными символами
                    result_paragraph = row_cells[4].paragraphs[0]
                    result_paragraph.clear()
                    if result['is_correct']:
                        run = result_paragraph.add_run("✓ Правильно")
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый цвет
                    else:
                        run = result_paragraph.add_run("✗ Неправильно")
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет
                
                # Добавляем разрыв страницы сразу после таблицы результатов
                doc.add_page_break()
                
                doc.save(word_path)
                log.info(f"Створено Word документ з результатами перевірки (fallback): {word_path}")
                return word_path
                
            except Exception as fallback_error:
                log.error(f"Помилка при створенні Word звіту (fallback): {fallback_error}")
                raise
        else:
            log.error(f"Помилка при створенні Word звіту: {e}")
            raise


def format_option_value(value):
    """Форматирует значение варианта ответа, убирая дробные части для целых чисел"""
    if isinstance(value, (int, float)):
        try:
            # Проверяем, является ли число целым
            if float(value) == int(float(value)):
                return str(int(float(value)))
            else:
                return f"{float(value):.10g}".replace('.', ',')
        except (ValueError, TypeError):
            return str(value).strip()
    else:
        # Для строковых значений проверяем, является ли это украинской буквой
        str_value = str(value).strip()
        ukrainian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ю', 'Я']
        if str_value.upper() in ukrainian_letters:
            return str_value.upper()
        # Пытаемся преобразовать в число только если это не украинская буква
        try:
            num_value = float(str_value.replace(',', '.'))
            if num_value == int(num_value):
                return str(int(num_value))
            else:
                return f"{num_value:.10g}".replace('.', ',')
        except (ValueError, TypeError):
            return str_value

def create_test_word(variants: List[Dict[str, Any]], output_dir: str, columns: int = 1, input_file_name: str = "", answer_format: str = "list", space_optimization: bool = False, test_class: str = "", test_date: str = "") -> str:
    """Создать Word документ с тестами для всех вариантов
    
    Args:
        variants: Список вариантов тестов
        output_dir: Папка для сохранения файлов
        columns: Количество колонок для размещения вопросов (всегда 1)
        input_file_name: Имя входного файла
        answer_format: Формат вариантов ответов ('list' или 'table')
        space_optimization: Минимизировать переводы строк для экономии места
        test_class: Класс для отображения в заголовке (опционально)
        test_date: Дата теста для отображения в заголовке (опционально)
    """
    try:
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        if input_file_name:
            word_path = os.path.join(output_dir, f"{input_file_name}_тести_{timestamp}.docx")
        else:
            word_path = os.path.join(output_dir, f"tests_{timestamp}.docx")
        
        doc = Document()
        
        # Устанавливаем минимальные поля для экономии места
        section = doc.sections[0]
        section.top_margin = Inches(0.5)     # 1.27 см
        section.bottom_margin = Inches(0.5)  # 1.27 см
        section.left_margin = Inches(0.5)    # 1.27 см
        section.right_margin = Inches(0.5)   # 1.27 см
        
        for variant in variants:
            # Заголовок варианта - общий по центру перед колонками
            title_parts = ["Тест"]
            if test_class:
                title_parts.append(f"Клас: {test_class}")
            if test_date:
                title_parts.append(f"Дата: {test_date}")
            title_parts.append(f"Варіант {variant['variant_number']}")
            
            heading = doc.add_heading(" - ".join(title_parts), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Настройка отступов для заголовка в режиме экономии места
            if space_optimization:
                heading.paragraph_format.space_before = Inches(0)
                heading.paragraph_format.space_after = Inches(0.05)
                heading.paragraph_format.line_spacing = 1.0
            
            # Инструкция
            instruction = doc.add_paragraph("Інструкція: Оберіть правильну відповідь для кожного питання.")
            instruction.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Настройка отступов для инструкции в режиме экономии места
            if space_optimization:
                instruction.paragraph_format.space_before = Inches(0)
                instruction.paragraph_format.space_after = Inches(0.02)
                instruction.paragraph_format.line_spacing = 1.0
            
            if not space_optimization:
                doc.add_paragraph()  # Пустая строка
            
            # Вычисляем общий вес всех заданий для нормализации баллов
            total_weight = sum(q['weight'] for q in variant['questions'])
            total_points = 12  # Общее количество баллов за тест
            
            # Вопросы в одноколоночной компоновке
            for i, question in enumerate(variant['questions'], 1):
                # Вычисляем баллы за это задание
                question_points = (question['weight'] / total_weight) * total_points
                # Форматируем баллы красиво
                if question_points == int(question_points):
                    points_str = f"({int(question_points)} балів)"
                else:
                    points_str = f"({format_number_with_comma(question_points, 1)} балів)"
                
                # Додаємо індикатор типу завдання
                task_type_indicator = get_task_type_indicator(question)
                points_and_type_str = f"{points_str} {task_type_indicator}"
                
                if space_optimization:
                    # В компактном режиме - номер, баллы и текст вопроса в одной строке
                    question_para = doc.add_paragraph()
                    # Добавляем номер и баллы жирным шрифтом
                    run1 = question_para.add_run(f"{i}. {points_and_type_str} ")
                    run1.bold = True
                    # Добавляем текст вопроса с обработкой математических формул
                    add_formatted_text_to_paragraph(question_para, question['question_text'])
                    question_para.style = 'Normal'
                    # Минимальные отступы для вопросов
                    question_para.paragraph_format.space_before = Inches(0)
                    question_para.paragraph_format.space_after = Inches(0.02)
                    question_para.paragraph_format.line_spacing = 1.0
                else:
                    # В обычном режиме - номер и баллы отдельно от текста вопроса
                    question_header = doc.add_paragraph(f"{i}. {points_and_type_str}")
                    question_header.runs[0].bold = True
                    
                    # Текст вопроса отдельной строкой с обработкой математических формул
                    question_para = doc.add_paragraph()
                    add_formatted_text_to_paragraph(question_para, question['question_text'])
                    question_para.style = 'Normal'
                
                # Варианты ответов в зависимости от типа задания
                if question['is_test_question']:
                    # Тестовое задание с вариантами ответов
                    if answer_format == 'table':
                        # Табличный формат - варианты ответов в таблице по ширине страницы
                        options = question['options']
                        num_options = len(options)
                        
                        # Создаем таблицу с одной строкой и количеством колонок равным количеству вариантов
                        table = doc.add_table(rows=1, cols=num_options)
                        table.style = 'Table Grid'
                        
                        # Растягиваем таблицу по всей ширине страницы
                        table.autofit = False
                        for col_idx, col in enumerate(table.columns):
                            col.width = Inches(6.5 / num_options)  # Равномерно распределяем по ширине
                        
                        # Додаємо світло-сіру рамку до таблиці
                        for row in table.rows:
                            for cell in row.cells:
                                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/><w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/></w:tcBorders>'))
                        
                        # Заполняем ячейки вариантами ответов
                        cells = table.rows[0].cells
                        # Українські літери без Ґ, Є, І, Ї, Й, Ь
                        ukrainian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ю', 'Я']
                        for j, option in enumerate(options):
                            # Форматируем числовые значения правильно
                            formatted_option = format_option_value(option)
                            # Очищаем ячейку и добавляем форматированный текст
                            cells[j].text = ""
                            para = cells[j].paragraphs[0]
                            letter = ukrainian_letters[j] if j < len(ukrainian_letters) else str(j + 1)
                            add_formatted_text_to_paragraph(para, f"{letter}) {formatted_option}")
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # Списочный формат - обычные варианты ответов
                        # Українські літери без Ґ, Є, І, Ї, Й, Ь
                        ukrainian_letters = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ю', 'Я']
                        for j, option in enumerate(question['options']):
                            # Форматируем числовые значения правильно
                            formatted_option = format_option_value(option)
                            option_para = doc.add_paragraph()
                            letter = ukrainian_letters[j] if j < len(ukrainian_letters) else str(j + 1)
                            add_formatted_text_to_paragraph(option_para, f"   {letter}) {formatted_option}")
                            option_para.style = 'Normal'
                            # Минимальные отступы для вариантов ответов в режиме экономии места
                            if space_optimization:
                                option_para.paragraph_format.space_before = Inches(0)
                                option_para.paragraph_format.space_after = Inches(0.01)
                                option_para.paragraph_format.line_spacing = 1.0
                else:
                    # Нетестовое задание - место для ответа
                    answer_para = doc.add_paragraph("Відповідь: ___________________________")
                    answer_para.style = 'Normal'
                    # Минимальные отступы для поля ответа в режиме экономии места
                    if space_optimization:
                        answer_para.paragraph_format.space_before = Inches(0)
                        answer_para.paragraph_format.space_after = Inches(0.02)
                        answer_para.paragraph_format.line_spacing = 1.0
                
                if not space_optimization:
                    doc.add_paragraph()  # Пустая строка между вопросами
            
            # Рядки для відповідей - по 6 завдань у рядку
            answers_paragraph = doc.add_paragraph()
            answers_run = answers_paragraph.add_run("Відповіді:")
            answers_run.bold = True
            # Мінімальні відступи для заголовка відповідей в режимі економії місця
            if space_optimization:
                answers_paragraph.paragraph_format.space_before = Inches(0.02)
                answers_paragraph.paragraph_format.space_after = Inches(0.01)
                answers_paragraph.paragraph_format.line_spacing = 1.0
            
            total_questions = len(variant['questions'])
            questions_per_row = 6  # По 6 завдань у рядку
            num_rows = (total_questions + questions_per_row - 1) // questions_per_row  # Округлення вгору
            
            current_q = 0
            for row_idx in range(num_rows):
                # Створюємо рядок з номерами завдань та місцями для відповідей
                row_text = ""
                for i in range(questions_per_row):
                    if current_q < total_questions:
                        # Резервуємо два символи для номера завдання
                        question_num = f"{current_q + 1:2d}"
                        row_text += f"{question_num}. _________________ "
                        current_q += 1
                    else:
                        # Додаємо пробіли для вирівнювання, якщо завдань менше 6
                        row_text += "                                   "
                
                # Додаємо параграф з рядком відповідей
                answer_para = doc.add_paragraph(row_text.rstrip())
                answer_para.style = 'Normal'
                # Мінімальні відступи для рядків відповідей в режимі економії місця
                if space_optimization:
                    answer_para.paragraph_format.space_before = Inches(0)
                    answer_para.paragraph_format.space_after = Inches(0.01)
                    answer_para.paragraph_format.line_spacing = 1.0
                
            
            # Разрыв страницы между вариантами (кроме последнего)
            if variant != variants[-1]:
                page_break = doc.add_page_break()
                if space_optimization:
                    # Мінімальні відступи для розриву сторінки в режимі економії місця
                    page_break_format = page_break._element.getparent()
                    if page_break_format is not None:
                        page_break_format.set(qn('w:before'), '0')
                        page_break_format.set(qn('w:after'), '0')
        
        doc.save(word_path)
        log.info(f"Word документ з тестами створено: {word_path}")
        return word_path
        
    except Exception as e:
        log.error(f"Помилка при створенні Word документа: {e}")
        raise


def read_test_word(file_path: str) -> pd.DataFrame:
    """Прочитать тест из Word документа и преобразовать в DataFrame"""
    try:
        doc = Document(file_path)
        questions_data = []
        
        current_question = None
        current_options = []
        question_number = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Проверяем, является ли это вопросом (начинается с числа и точки)
            if text and text[0].isdigit() and '. ' in text:
                # Сохраняем предыдущий вопрос
                if current_question and current_options:
                    questions_data.append({
                        'question': current_question,
                        'option_1': current_options[0] if len(current_options) > 0 else '',
                        'option_2': current_options[1] if len(current_options) > 1 else '',
                        'option_3': current_options[2] if len(current_options) > 2 else '',
                        'option_4': current_options[3] if len(current_options) > 3 else '',
                        'correct_answer': 1  # По умолчанию первый вариант
                    })
                
                # Начинаем новый вопрос
                question_number += 1
                current_question = text.split('. ', 1)[1] if '. ' in text else text
                current_options = []
            
            # Проверяем, является ли это вариантом ответа
            elif text and (text.startswith('   1)') or text.startswith('   2)') or 
                          text.startswith('   3)') or text.startswith('   4)')):
                option_text = text.split(') ', 1)[1] if ') ' in text else text
                current_options.append(option_text)
        
        # Сохраняем последний вопрос
        if current_question and current_options:
            questions_data.append({
                'question': current_question,
                'option_1': current_options[0] if len(current_options) > 0 else '',
                'option_2': current_options[1] if len(current_options) > 1 else '',
                'option_3': current_options[2] if len(current_options) > 2 else '',
                'option_4': current_options[3] if len(current_options) > 3 else '',
                'correct_answer': 1  # По умолчанию первый вариант
            })
        
        if not questions_data:
            raise ValueError("Не вдалося знайти питання в Word документі")
        
        df = pd.DataFrame(questions_data)
        
        # Додаємо номери питань
        df['question_number'] = range(1, len(df) + 1)
        
        # Додаємо вагу питань (за замовчуванням 1)
        df['weight'] = 1
        
        # Рахуємо кількість варіантів відповідей для кожного питання
        df['option_count'] = 0
        for idx, row in df.iterrows():
            count = 0
            for i in range(1, 5):  # option_1 до option_4
                if row[f'option_{i}'] and str(row[f'option_{i}']).strip():
                    count += 1
            df.at[idx, 'option_count'] = count
        
        # Конвертуємо option_ колонки в рядки для сумісності з pyarrow
        option_cols = [col for col in df.columns if col.startswith('option_')]
        for col in option_cols:
            df[col] = df[col].astype(str)
        
        # Визначаємо тип питання (тестове чи відкрите)
        df['is_test_question'] = (df['option_count'] >= 2) & df['correct_answer'].notna() & (df['correct_answer'] != 'nan')
        
        # Приводимо стовпець correct_answer до object типу для уникнення попереджень
        df['correct_answer'] = df['correct_answer'].astype('object')
        
        # Правильні відповіді для тестових завдань вже оброблені вище (рядки 330-370)
        
        # Для відкритих завдань зберігаємо правильну відповідь як текст
        open_mask = ~df['is_test_question']
        if open_mask.any():
            for idx in df[open_mask].index:
                df.at[idx, 'correct_answer'] = str(df.at[idx, 'correct_answer'])
        
        log.info(f"З Word документа завантажено {len(df)} питань")
        return df
        
    except Exception as e:
        log.error(f"Помилка при читанні Word документа: {e}")
        raise


def export_answers_to_word(variants: List[Dict[str, Any]], output_dir: str, input_file_name: str = "", test_class: str = "", test_date: str = "") -> str:
    """Экспортировать ответы всех вариантов в Word документ
    
    Args:
        variants: Список вариантов тестов
        output_dir: Папка для сохранения файлов
        input_file_name: Имя входного файла
        test_class: Класс для отображения в заголовке (опционально)
        test_date: Дата теста для отображения в заголовке (опционально)
    """
    try:
        os.makedirs(output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        if input_file_name:
            word_path = os.path.join(output_dir, f"{input_file_name}_Відповіді_{timestamp}.docx")
        else:
            word_path = os.path.join(output_dir, f"answers_{timestamp}.docx")
        
        doc = Document()
        
        # Устанавливаем минимальные поля для экономии места
        section = doc.sections[0]
        section.top_margin = Inches(0.5)     # 1.27 см
        section.bottom_margin = Inches(0.5)  # 1.27 см
        section.left_margin = Inches(0.5)    # 1.27 см
        section.right_margin = Inches(0.5)   # 1.27 см
        
        # Заголовок
        title_parts = ["Відповіді до тестів"]
        if test_class:
            title_parts.append(f"Клас: {test_class}")
        if test_date:
            title_parts.append(f"Дата: {test_date}")
        
        heading = doc.add_heading(" - ".join(title_parts), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Инструкция
        instruction = doc.add_paragraph(
            "Цей документ містить правильні відповіді для всіх варіантів тестів. "
            "Використовуйте його для перевірки робіт учнів."
        )
        instruction.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()  # Пустая строка
        
        # Ответы для каждого варианта
        for variant in variants:
            # Заголовок варианта
            variant_heading = doc.add_heading(f'Варіант {variant["variant_number"]}', level=2)
            variant_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Ответы в виде таблицы
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Номер питання'
            hdr_cells[1].text = 'Правильна відповідь'
            
            # Делаем заголовки жирными
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Добавляем строки с ответами
            for i, answer in enumerate(variant['answer_key'], 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = str(answer)
            
            # Також добавляємо ответы в строку для удобства
            answers_line = doc.add_paragraph()
            answers_line.add_run('Відповіді в рядок: ').bold = True
            answers_line.add_run(', '.join([str(ans) for ans in variant['answer_key']]))
            
            doc.add_paragraph()  # Пустая строка между вариантами
        
        doc.save(word_path)
        log.info(f"Word документ з відповідями створено: {word_path}")
        return word_path
        
    except Exception as e:
        log.error(f"Помилка при експорті відповідей в Word: {e}")
        raise

def generate_test_template(output_dir: str) -> str:
    """
    Генерує Excel шаблон для тестів.
    
    Args:
        output_dir: Директорія для збереження шаблону
        
    Returns:
        Шлях до створеного шаблону
    """
    try:
        # Створюємо директорію якщо не існує
        os.makedirs(output_dir, exist_ok=True)
        
        # Шлях до шаблону
        test_template_path = os.path.join(output_dir, "Шаблон_тесту.xlsx")
        
        # Створюємо шаблон
        create_test_template(test_template_path)
        
        log.info(f"Створено шаблон: {test_template_path}")
        
        return test_template_path
        
    except Exception as e:
        log.error(f"Помилка при створенні шаблону: {e}")
        raise

def generate_neural_query_document(output_dir: str) -> str:
    """
    Генерирует документ Word с запросом для нейросети.
    
    Args:
        output_dir: Директория для сохранения документа
        
    Returns:
        str: Путь к созданному документу
    """
    try:
        # Создаем директорию если не существует
        os.makedirs(output_dir, exist_ok=True)
        
        # Путь для документа
        query_doc_path = os.path.join(output_dir, "Запрос_для_нейросети.docx")
        
        # Создаем документ
        create_neural_query_document(query_doc_path)
        
        log.info(f"Документ з запитом для нейромережі створено: {query_doc_path}")
        
        return query_doc_path
        
    except Exception as e:
        log.error(f"Помилка при створенні документа з запитом: {e}")
        raise


def save_student_result_to_excel(check_result: Dict[str, Any], student_info: Dict[str, str], 
                                work_name: str, excel_file_path: str, key_file_name: str = "") -> None:
    """
    Зберігає результат перевірки учня в Excel файл.
    
    Args:
        check_result: Результати перевірки
        student_info: Інформація про учня (клас, ПІБ)
        work_name: Назва роботи
        excel_file_path: Шлях до Excel файлу
        key_file_name: Назва файлу-ключа
    """
    try:
        # Підготовка даних для збереження
        check_date = datetime.now().strftime("%Y-%m-%d %H:%M")
        variant = check_result['variant_number']
        student_class = student_info.get('class', '')
        student_name = student_info.get('full_name', '')
        
        # Детальні бали за кожне завдання (чисті числа)
        detailed_points = []
        detailed_max_points = []
        for result in check_result['detailed_results']:
            detailed_points.append(result['points'])
            detailed_max_points.append(result['max_points'])
        
        # Загальні результати (чисті числа)
        total_percentage = check_result['score_percentage']
        weighted_score = check_result['weighted_score']
        max_score = check_result['max_score']
        
        # Створюємо рядок даних з чистими числами
        row_data = {
            'Файл-ключ': key_file_name,
            'Дата перевірки': check_date,
            'Назва роботи': work_name,
            'Клас': student_class,
            'Учень': student_name,
            'Варіант': variant,
            'Загальний відсоток (%)': total_percentage,
            'Отримано балів': weighted_score,
            'Максимум балів': max_score
        }
        
        # Додаємо детальні бали за кожне завдання (чисті числа)
        for i, (points, max_points) in enumerate(zip(detailed_points, detailed_max_points), 1):
            row_data[f'Завдання {i} (бали)'] = points
            row_data[f'Завдання {i} (макс)'] = max_points
        
        # Перевіряємо, чи існує файл
        if os.path.exists(excel_file_path):
            # Читаємо існуючий файл
            try:
                existing_df = pd.read_excel(excel_file_path)
            except Exception:
                # Якщо файл пошкоджений, створюємо новий
                existing_df = pd.DataFrame()
        else:
            # Створюємо новий DataFrame
            existing_df = pd.DataFrame()
        
        # Створюємо новий рядок
        new_row_df = pd.DataFrame([row_data])
        
        # Об'єднуємо з існуючими даними
        if not existing_df.empty:
            # Переконуємося, що всі колонки присутні в обох DataFrame
            all_columns = list(set(existing_df.columns.tolist() + new_row_df.columns.tolist()))
            
            # Переупорядковуємо колонки: основні спочатку, потім завдання
            base_columns = ['Файл-ключ', 'Дата перевірки', 'Назва роботи', 'Клас', 'Учень', 'Варіант', 
                          'Загальний відсоток (%)', 'Отримано балів', 'Максимум балів']
            task_columns = [col for col in all_columns if col.startswith('Завдання')]
            task_columns.sort(key=lambda x: (int(x.split()[1]) if x.split()[1].isdigit() else 0, x))
            
            ordered_columns = base_columns + task_columns
            
            # Додаємо відсутні колонки
            for col in ordered_columns:
                if col not in existing_df.columns:
                    existing_df[col] = ''
                if col not in new_row_df.columns:
                    new_row_df[col] = ''
            
            # Переупорядковуємо колонки
            existing_df = existing_df[ordered_columns]
            new_row_df = new_row_df[ordered_columns]
            
            # Об'єднуємо
            result_df = pd.concat([existing_df, new_row_df], ignore_index=True)
        else:
            result_df = new_row_df
        
        # Створюємо папку, якщо вона не існує
        os.makedirs(os.path.dirname(excel_file_path), exist_ok=True)
        
        # Зберігаємо в Excel
        result_df.to_excel(excel_file_path, index=False)
        
        log.info(f"Результат учня збережено в Excel файл: {excel_file_path}")
        
    except Exception as e:
        log.error(f"Помилка при збереженні результату в Excel: {e}")
        raise
